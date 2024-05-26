import streamlit as st
import pandas as pd
import spacy
from spacy.tokens import Span
from spacy.language import Language
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import matplotlib.dates as mpl_dates
import numpy as np
import dataframe_image as dfi
from docx import Document
from docx.shared import Inches
from textwrap import wrap
from io import BytesIO
import openai
import os

# Function to get the OpenAI API key
def get_api_key():
    if 'openai' in st.secrets:
        return st.secrets['openai']['api_key']
    return os.getenv('OPENAI_API_KEY', 'Your-OpenAI-API-Key')

openai.api_key = get_api_key()  # Set the OpenAI API key

# Download spaCy model if not already present
spacy_model = "en_core_web_sm"
try:
    nlp = spacy.load(spacy_model)
except OSError:
    from spacy.cli import download
    download(spacy_model)
    nlp = spacy.load(spacy_model)

# Add custom rule to recognize vessel names as ORGANIZATION entities
@Language.component("vessel_name_rule")
def vessel_name_rule(doc):
    pattern = [{'LOWER': {'IN': ['vessel', 'ship', 'ss', 'mv']}},
               {'IS_ALPHA': True, 'OP': '+'}]
    matcher = spacy.matcher.Matcher(nlp.vocab)
    matcher.add('VESSEL_NAME', [pattern])
    matches = matcher(doc)
    spans = [Span(doc, start, end, label="ORG") for match_id, start, end in matches]
    doc.ents = list(doc.ents) + spans
    return doc

# Define the PhraseMatcher outside the component
phrase_patterns = [
    "hull condition", "condition of hull", "hull performance", "performance of hull",
    "propeller performance", "condition of propeller", "hull and propeller condition",
    "performance of hull and propeller", "propeller condition", "performance of propeller",
    "vessel performance report", "performance report"
]
matcher = PhraseMatcher(nlp.vocab)
patterns = [nlp(text) for text in phrase_patterns]
matcher.add("CONDITION", patterns)

# Create a custom component for PhraseMatcher
@Language.component("phrase_matcher")
def phrase_matcher(doc):
    matches = matcher(doc)
    spans = [Span(doc, start, end, label="CONDITION") for match_id, start, end in matches]
    doc.ents = list(doc.ents) + spans
    return doc

# Add custom components to the pipeline
nlp.add_pipe("vessel_name_rule", after='ner')
nlp.add_pipe("phrase_matcher", last=True)

# Function to extract vessel name from the query
def extract_vessel_name(query):
    doc = nlp(query)
    for ent in doc.ents:
        if ent.label_ == "ORG":  # Assuming vessel names are recognized as ORG
            return ent.text.lower()  # Return in lowercase for consistency
    return None

# Function to detect if the query is about hull or propeller conditions
def is_hull_or_propeller_query(query):
    doc = nlp(query)
    for ent in doc.ents:
        if ent.label_ == "CONDITION":
            return True
    return False

def load_hull_data(file_url):
    df = pd.read_csv(file_url)
    df['REPORT_DATE'] = pd.to_datetime(df['REPORT_DATE'], errors='coerce')
    df.set_index('REPORT_DATE', inplace=True)
    df['VESSEL_NAME'] = df['VESSEL_NAME'].str.lower()
    return df

def handle_hull_performance_query(df, vessel_name):
    column_name = "HULL_ROUGHNESS_POWER_LOSS"
    three_months_ago = df.index.max() - pd.DateOffset(months=3)
    filtered_df = df[(df['VESSEL_NAME'] == vessel_name) & (df.index >= three_months_ago)]
    
    if filtered_df.empty or filtered_df[column_name].dropna().empty:
        return f"No valid data available for the vessel {vessel_name.upper()} in the last three months.", None
    
    average_power_loss = filtered_df[column_name].dropna().mean()
    if average_power_loss < 15:
        hull_condition = "Good"
        recommendation = "The hull condition of this vessel is Good."
        condition_color = "green"
    elif 15 <= average_power_loss <= 25:
        hull_condition = "Average"
        recommendation = "Hull cleaning and propeller polishing is recommended at the next convenient/economical opportunity."
        condition_color = "orange"
    else:
        hull_condition = "Poor"
        recommendation = "Hull cleaning and propeller polishing is recommended at the earliest."
        condition_color = "red"
    
    potential_fuel_savings = filtered_df['HULL_EXCESS_FUEL_OIL_MTD'].dropna().mean()
    result_data = {
        "vessel_name": vessel_name.upper(),
        "excess_power_percentage": average_power_loss,
        "hull_condition": hull_condition,
        "condition_color": condition_color,
        "potential_fuel_savings": potential_fuel_savings,
        "recommendation": recommendation,
        "filtered_df": filtered_df
    }
    
    return None, result_data

def load_performance_data(file_url):
    df = pd.read_csv(file_url)
    df['REPORTDATE'] = pd.to_datetime(df['REPORTDATE'], errors='coerce')
    df.set_index('REPORTDATE', inplace=True)
    df['VESSEL_NAME'] = df['VESSEL_NAME'].str.lower()
    return df

def handle_vessel_performance_query(df, vessel_name):
    one_month_ago = df.index.max() - pd.DateOffset(months=1)
    filtered_df = df[(df['VESSEL_NAME'] == vessel_name) & (df.index >= one_month_ago)]
    
    if filtered_df.empty or filtered_df['ME_SFOC'].dropna().empty or filtered_df['ACTUAL_ME_POWER'].dropna().empty:
        return f"No valid data available for the vessel {vessel_name.upper()} in the last one month.", None
    
    avg_sfoc = filtered_df['ME_SFOC'].dropna().mean()
    avg_power = filtered_df['ACTUAL_ME_POWER'].dropna().mean()
    
    if avg_sfoc < 160 or avg_sfoc > 250:
        condition = "Anomalous"
        recommendation = "Please check the reported power/ME Fuel consumption."
        condition_color = "yellow"
    elif 160 <= avg_sfoc < 190:
        condition = "Good"
        recommendation = "-"
        condition_color = "green"
    elif 190 <= avg_sfoc < 210:
        condition = "Average"
        recommendation = "Analyse ME performance report and carry out adjustments/maintenance as required."
        condition_color = "orange"
    else:
        condition = "Poor"
        recommendation = "Urgently analyse the performance and carry out maintenance as required."
        condition_color = "red"
    
    fuel_saving_potential = avg_sfoc * avg_power * 24 / (10**6)
    result_data = {
        "vessel_name": vessel_name.upper(),
        "avg_sfoc": avg_sfoc,
        "condition": condition,
        "condition_color": condition_color,
        "fuel_saving_potential": fuel_saving_potential,
        "recommendation": recommendation,
        "filtered_df": filtered_df
    }
    
    return None, result_data

def apply_background_color(row):
    if row['Condition'] == 'Good':
        return ['background-color: green' if i == 2 else '' for i in range(len(row))]
    elif row['Condition'] == 'Average':
        return ['background-color: orange' if i == 2 else '' for i in range(len(row))]
    elif row['Condition'] == 'Poor':
        return ['background-color: red' if i == 2 else '' for i in range(len(row))]
    elif row['Condition'] == 'Anomalous':
        return ['background-color: yellow' if i == 2 else '' for i in range(len(row))]
    else:
        return ['background-color: white' if i == 2 else '' for i in range(len(row))]

def generate_table_image(df, filename):
    dfi.export(df, filename, table_conversion='matplotlib')

def save_document_with_tables_and_charts(hull_result_data, performance_result_data, query, filename='performance_report.docx'):
    document = Document()
    
    # Hull Performance Section
    if 'hull' in query or 'vessel performance' in query:
        document.add_heading('Hull Performance Report', level=1)
        if hull_result_data:
            if os.path.exists('hull_performance_table.png'):
                document.add_picture('hull_performance_table.png', width=Inches(6))

            for i, result in enumerate(hull_result_data):
                if isinstance(result['excess_power_percentage'], (int, float)):
                    document.add_heading(f"Vessel: {result['vessel_name']}", level=2)
                    document.add_paragraph(f"Excess Power Percentage: {result['excess_power_percentage']:.2f}%")
                    document.add_paragraph(f"Hull Condition: {result['hull_condition']}")
                    document.add_paragraph(f"Potential Fuel Savings: {result['potential_fuel_savings']:.2f} mt/d")
                    document.add_paragraph(f"Recommendation: {result['recommendation']}")
                    
                    # Add chart
                    chart_stream = BytesIO()
                    plt.figure(figsize=(12, 8))
                    filtered_df = result['filtered_df']
                    x = filtered_df.index
                    y = filtered_df['HULL_ROUGHNESS_POWER_LOSS'].dropna()
                    x = x[filtered_df['HULL_ROUGHNESS_POWER_LOSS'].notna()]
                    plt.scatter(x, y, color='white')
                    z = np.polyfit(mpl_dates.date2num(x), y, 1)
                    p = np.poly1d(z)
                    plt.plot(x, p(mpl_dates.date2num(x)), color='red')
                    plt.title(f"Excess Power % (compared with baseline) - {result['vessel_name']}", color='white')
                    plt.xlabel('Date', color='white')
                    plt.ylabel('Excess Power %', color='white')
                    plt.gca().set_facecolor('#000C20')
                    plt.gcf().set_facecolor('#000C20')
                    plt.xticks(color='white')
                    plt.yticks(color='white')
                    plt.grid(True, color='gray', linestyle='--', linewidth=0.5)
                    plt.savefig(chart_stream, format='png', dpi=300, bbox_inches='tight')
                    chart_stream.seek(0)
                    document.add_picture(chart_stream, width=Inches(6))
                    plt.close()
                    
                    if i < len(hull_result_data) - 1:
                        document.add_page_break()
        else:
            document.add_paragraph("No valid data available for hull performance.")

    # Vessel Performance Section
    if 'vessel performance' in query:
        document.add_heading('Machinery Performance Report', level=1)
        if performance_result_data:
            if os.path.exists('vessel_performance_table.png'):
                document.add_picture('vessel_performance_table.png', width=Inches(6))
            
            for result in performance_result_data:
                document.add_heading(f"Vessel: {result['vessel_name']}", level=2)
                document.add_paragraph(f"ME SFOC: {result['avg_sfoc']:.2f}")
                document.add_paragraph(f"Condition: {result['condition']}")
                document.add_paragraph(f"Fuel Saving Potential: {result['fuel_saving_potential']:.2f} mt/d")
                document.add_paragraph(f"Recommendation: {result['recommendation']}")
                
                # Add chart
                chart_stream = BytesIO()
                plt.figure(figsize=(12, 8))
                df = result['filtered_df']
                x = df.index
                y = df['ME_SFOC'].dropna()
                x = x[df['ME_SFOC'].notna()]
                plt.plot(x, y, color='blue', marker='o', linestyle='-')
                plt.title(f"ME SFOC - {result['vessel_name']}")
                plt.xlabel('Date')
                plt.ylabel('ME SFOC')
                plt.grid(True)
                plt.savefig(chart_stream, format='png', dpi=300, bbox_inches='tight')
                chart_stream.seek(0)
                document.add_picture(chart_stream, width=Inches(6))
                plt.close()
        else:
            document.add_paragraph("No valid data available for machinery performance.")

    document.save(filename)
    return filename

def display_table(df):
    plt.figure(figsize=(12, len(df) * 0.5 + 1))
    plt.axis('off')
    plt.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center', edges='B')
    plt.tight_layout()
    plt.show()

def display_chart(df, column_name, title):
    plt.figure(figsize=(12, 6))
    df[column_name].plot(kind='line')
    plt.title(title)
    plt.xlabel('Date')
    plt.ylabel(column_name)
    plt.grid(True)
    plt.show()

# Streamlit app
def main():
    st.title("Vessel Performance Report Generator")

    # Use the provided downloadable URLs
    hull_file_url = 'https://drive.google.com/uc?export=download&id=1dNjLhrvUIWuwiQh9QIZ-7Xc0ba3ze23m'
    performance_file_url = 'https://drive.google.com/uc?export=download&id=1dK9C9niJAm040YCOij-UxziBZIfkJTij'
    
    st.write("Debug: Streamlit app is running.")

    vessel_names = st.text_input("Enter vessel names (comma-separated)").lower().split(', ')
    st.write(f"Debug: Vessel names - {vessel_names}")

    query = st.text_input("Enter your query (e.g., 'hull performance', 'vessel performance')").lower()
    st.write(f"Debug: Query - {query}")

    if st.button("Generate Report"):
        st.write("Debug: Button clicked.")
        if hull_file_url and performance_file_url and vessel_names and query:
            st.write("Debug: Processing inputs.")
            df_hull = load_hull_data(hull_file_url)
            df_performance = load_performance_data(performance_file_url)

            hull_result_data = []
            performance_result_data = []

            if 'hull' in query or 'vessel performance' in query:
                for vessel_name in vessel_names:
                    error, result = handle_hull_performance_query(df_hull, vessel_name)
                    if result:
                        hull_result_data.append(result)
                if hull_result_data:
                    hull_table_data = []
                    for result in hull_result_data:
                        hull_table_data.append([
                            result["vessel_name"],
                            f"{result['excess_power_percentage']:.2f}%" if isinstance(result['excess_power_percentage'], (int, float)) else "-",
                            result["hull_condition"],
                            f"{result['potential_fuel_savings']:.2f}" if isinstance(result['potential_fuel_savings'], (int, float)) else "-",
                            result["recommendation"]
                        ])
                    hull_table_df = pd.DataFrame(hull_table_data, columns=["Vessel Name", "Excess Power %", "Condition", "Potential Fuel Savings (mt/d)", "Recommendation"])
                    hull_table_df['Recommendation'] = hull_table_df['Recommendation'].apply(lambda x: '\n'.join(wrap(x, 40)))
                    generate_table_image(hull_table_df.style.apply(apply_background_color, axis=1), 'hull_performance_table.png')
                    st.write("Hull performance data processed and table generated.")
                    st.dataframe(hull_table_df)
                else:
                    st.write("No hull performance data found for the given vessels.")

            if 'vessel performance' in query:
                for vessel_name in vessel_names:
                    error, result = handle_vessel_performance_query(df_performance, vessel_name)
                    if result:
                        performance_result_data.append(result)
                if performance_result_data:
                    performance_table_data = []
                    for result in performance_result_data:
                        performance_table_data.append([
                            result["vessel_name"],
                            f"{result['avg_sfoc']:.2f}" if isinstance(result['avg_sfoc'], (int, float)) else "-",
                            result["condition"],
                            f"{result['fuel_saving_potential']:.2f}" if isinstance(result['fuel_saving_potential'], (int, float)) else "-",
                            result["recommendation"]
                        ])
                    performance_table_df = pd.DataFrame(performance_table_data, columns=["Vessel Name", "ME SFOC", "Condition", "Fuel Saving Potential (mt/d)", "Recommendation"])
                    performance_table_df['Recommendation'] = performance_table_df['Recommendation'].apply(lambda x: '\n'.join(wrap(x, 40)))
                    generate_table_image(performance_table_df.style.apply(apply_background_color, axis=1), 'vessel_performance_table.png')
                    st.write("Vessel performance data processed and table generated.")
                    st.dataframe(performance_table_df)
                else:
                    st.write("No vessel performance data found for the given vessels.")

            # Save document with tables and charts
            if 'hull' in query or 'vessel performance' in query:
                report_path = save_document_with_tables_and_charts(hull_result_data, performance_result_data, query, filename='performance_report.docx')
                with open(report_path, "rb") as file:
                    btn = st.download_button(label="Download Report", data=file, file_name="performance_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.write(f"Report generated and saved at {report_path}.")

if __name__ == "__main__":
    main()
